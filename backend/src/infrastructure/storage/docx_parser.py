"""Dịch vụ phân tích nội dung tệp văn bản Word (.docx)."""

import io

from docx import Document
from src.common.exceptions.error_codes import ErrorCode
from src.common.exceptions.infrastructure import InfrastructureException
from src.domain.data_source.file_parser import ParsedRequirementResult


class DocxParser:
    """Công cụ đọc và bóc tách nội dung văn bản từ tệp Word (.docx)."""

    def parse(self, file_bytes: bytes, filename: str, saved_path: str = "") -> ParsedRequirementResult:
        """Đọc tệp .docx và trích xuất nội dung văn bản trực tiếp từ bộ nhớ.

        Args:
            file_bytes: Dữ liệu nhị phân của tệp Word.
            filename: Tên tệp gốc.
            saved_path: Đường dẫn lưu trữ tệp trên server (mặc định là rỗng).

        Returns:
            ParsedRequirementResult: Nội dung văn bản yêu cầu bóc tách được.
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            content_parts: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    content_parts.append(table_text)

            full_text = "\n\n".join(content_parts)
            return ParsedRequirementResult(
                name=filename,
                file_path=saved_path,
                content=full_text,
            )
        except Exception as exc:
            raise InfrastructureException(
                code=ErrorCode.FILE_PARSING_ERROR,
                message=f"Không thể đọc nội dung tệp Word: {filename}",
            ) from exc

    def _extract_table_text(self, table: object) -> str:
        """Trích xuất văn bản từ bảng trong tài liệu docx."""
        rows_text: list[str] = []
        for row in getattr(table, "rows", []):
            cells_text = [cell.text.strip() for cell in getattr(row, "cells", []) if cell.text.strip()]
            if cells_text:
                rows_text.append(" | ".join(cells_text))
        return "\n".join(rows_text)
