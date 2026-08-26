"""Parser TXT, Markdown và DOCX cho Requirement Documents."""

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from src.application.requirement_files.i_requirement_file_service import (
    IRequirementDocumentParser,
)
from src.common.exceptions.business import BusinessException
from src.common.exceptions.error_codes import ErrorCode
from src.domain.requirement_file.enums import RequirementFileType


class RequirementDocumentParser(IRequirementDocumentParser):
    """Ẩn chi tiết python-docx khỏi application boundary."""

    def parse(self, filename: str, content: bytes) -> tuple[RequirementFileType, str]:
        """Parse document và từ chối content rỗng hoặc không hợp lệ."""
        file_type = self._file_type(filename)
        try:
            text = self._extract(file_type, content).strip()
        except UnicodeDecodeError as exc:
            raise BusinessException(
                ErrorCode.FILE_PARSING_ERROR, "Requirement Document phải dùng UTF-8."
            ) from exc
        except (BadZipFile, PackageNotFoundError, KeyError, ValueError) as exc:
            raise BusinessException(
                ErrorCode.FILE_PARSING_ERROR, "Không thể parse Requirement Document."
            ) from exc
        if not text:
            raise BusinessException(
                ErrorCode.FILE_EMPTY, "Requirement Document không có nội dung text."
            )
        return file_type, text

    @staticmethod
    def _file_type(filename: str) -> RequirementFileType:
        extension = Path(filename.strip()).suffix.lower().removeprefix(".")
        try:
            return RequirementFileType(extension.upper())
        except ValueError as exc:
            raise BusinessException(
                ErrorCode.INVALID_FILE_FORMAT,
                "Requirement Document chỉ hỗ trợ DOCX, TXT hoặc MD.",
            ) from exc

    @staticmethod
    def _extract(file_type: RequirementFileType, content: bytes) -> str:
        if file_type != RequirementFileType.DOCX:
            return content.decode("utf-8")
        document = Document(BytesIO(content))
        paragraphs = [item.text for item in document.paragraphs if item.text.strip()]
        table_rows = (
            " | ".join(cell.text.strip() for cell in row.cells)
            for table in document.tables
            for row in table.rows
        )
        return "\n".join((*paragraphs, *(row for row in table_rows if row.strip(" |"))))
