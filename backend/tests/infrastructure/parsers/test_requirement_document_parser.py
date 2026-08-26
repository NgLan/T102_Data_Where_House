"""Requirement Document parser contracts."""

from io import BytesIO

import pytest
from docx import Document
from src.common.exceptions.business import BusinessException
from src.common.exceptions.error_codes import ErrorCode
from src.domain.requirement_file.enums import RequirementFileType
from src.infrastructure.parsers.requirement_document_parser import (
    RequirementDocumentParser,
)


@pytest.mark.parametrize(
    ("filename", "content", "file_type"),
    [
        ("requirements.txt", "Doanh thu theo tháng".encode(), RequirementFileType.TXT),
        ("requirements.MD", b"# Revenue\n\nBy month", RequirementFileType.MD),
    ],
)
def test_parse_utf8_text_formats(
    filename: str, content: bytes, file_type: RequirementFileType
) -> None:
    parsed_type, text = RequirementDocumentParser().parse(filename, content)

    assert parsed_type is file_type
    assert text


def test_parse_docx_includes_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Revenue reporting")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Net revenue"
    content = BytesIO()
    document.save(content)

    file_type, text = RequirementDocumentParser().parse("brief.docx", content.getvalue())

    assert file_type is RequirementFileType.DOCX
    assert text == "Revenue reporting\nMetric | Net revenue"


@pytest.mark.parametrize(
    ("filename", "content", "code"),
    [
        ("brief.pdf", b"text", ErrorCode.INVALID_FILE_FORMAT),
        ("brief.txt", b"\xff", ErrorCode.FILE_PARSING_ERROR),
        ("brief.txt", b"  \n", ErrorCode.FILE_EMPTY),
        ("brief.docx", b"not-a-docx", ErrorCode.FILE_PARSING_ERROR),
    ],
)
def test_reject_invalid_or_empty_documents(
    filename: str, content: bytes, code: ErrorCode
) -> None:
    with pytest.raises(BusinessException) as raised:
        RequirementDocumentParser().parse(filename, content)

    assert raised.value.code is code
