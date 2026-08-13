import os
import re
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_paragraph_with_runs(p, text):
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = p.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78)
        elif token.startswith("*") and token.endswith("*"):
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(38, 38, 38)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip("\r\n")

        # Code block check
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\r\n"))
                i += 1
            i += 1  # skip closing ```

            code_text = "\n".join(code_lines)

            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_background(cell, "F5F7FA")
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(40, 44, 52)

            doc.add_paragraph()
            continue

        # Markdown Table check
        if line.strip().startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            matrix = []
            for tline in table_lines:
                cells = [c.strip() for c in tline.split("|")[1:-1]]
                if all(re.match(r"^:?-+:?$", c) for c in cells):
                    continue
                matrix.append(cells)

            if matrix:
                num_rows = len(matrix)
                num_cols = max(len(row) for row in matrix)
                tbl = doc.add_table(rows=num_rows, cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                for r_idx, row in enumerate(matrix):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < num_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(0)
                            if r_idx == 0:
                                set_cell_background(cell, "1F497D")
                                add_paragraph_with_runs(p, cell_text)
                                for r in p.runs:
                                    r.font.color.rgb = RGBColor(255, 255, 255)
                                    r.bold = True
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, "F2F5F9")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                                add_paragraph_with_runs(p, cell_text)

                doc.add_paragraph()
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:])
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 73, 125)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(line[3:])
            run.font.name = "Calibri"
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:])
            run.font.name = "Calibri"
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(68, 114, 196)
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[5:])
            run.font.name = "Calibri"
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(68, 114, 196)
            i += 1
            continue

        # Horizontal Rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue

        # Bullet lists
        if line.strip().startswith(("- ", "* ")):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_with_runs(p, text)
            i += 1
            continue

        # Numbered lists
        m_num = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m_num:
            text = m_num.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_with_runs(p, text)
            i += 1
            continue

        # Normal paragraph
        if line.strip():
            p = doc.add_paragraph()
            add_paragraph_with_runs(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Successfully saved docx to: {docx_path}")


if __name__ == "__main__":
    import sys

    src_md = (
        sys.argv[1]
        if len(sys.argv) > 1
        else r"c:\T-102\P-102\docs\API_SPECIFICATION_FRONTEND_BACKEND.md"
    )
    dst_docx = (
        sys.argv[2]
        if len(sys.argv) > 2
        else r"c:\T-102\P-102\docs\API_SPECIFICATION_FRONTEND_BACKEND.docx"
    )
    convert_md_to_docx(src_md, dst_docx)
